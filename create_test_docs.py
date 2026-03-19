"""
Script to create sample test documents for testing the document comparison app.
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_test_doc1():
    """Create first test document."""
    doc = Document()

    # Add title
    title = doc.add_heading('Product Requirements Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Overview section
    doc.add_heading('1. Overview', 1)
    doc.add_paragraph(
        'This document outlines the requirements for the new mobile application. '
        'The app will provide users with real-time notifications and task management '
        'capabilities. Our goal is to create an intuitive, user-friendly interface '
        'that enhances productivity.'
    )

    # Features section
    doc.add_heading('2. Features', 1)
    doc.add_paragraph('The application will include the following key features:')

    features = [
        'User authentication with OAuth2 protocol',
        'Push notifications for task updates and reminders',
        'Calendar integration with Google Calendar and Outlook',
        'Collaborative task sharing with team members',
        'Real-time synchronization across devices'
    ]

    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    # Technical Specifications
    doc.add_heading('3. Technical Specifications', 1)
    doc.add_paragraph('The technology stack for this project:')

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Medium Grid 1 Accent 1'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology'
    hdr_cells[2].text = 'Version'

    # Data rows
    data = [
        ('Frontend', 'React Native', '0.72'),
        ('Backend', 'Node.js', '18.x'),
        ('Database', 'PostgreSQL', '15.0')
    ]

    for i, (comp, tech, ver) in enumerate(data, start=1):
        row = table.rows[i].cells
        row[0].text = comp
        row[1].text = tech
        row[2].text = ver

    # Timeline
    doc.add_heading('4. Timeline', 1)
    doc.add_paragraph(
        'Phase 1: Requirements gathering - 2 weeks\n'
        'Phase 2: Design and architecture - 3 weeks\n'
        'Phase 3: Development - 8 weeks\n'
        'Phase 4: Testing and QA - 2 weeks\n'
        'Phase 5: Deployment - 1 week'
    )

    doc.save('data/uploads/test_doc1.docx')
    print('✅ Created test_doc1.docx')


def create_test_doc2():
    """Create second test document (similar but with differences)."""
    doc = Document()

    # Add title
    title = doc.add_heading('Product Requirements Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Overview section (similar wording)
    doc.add_heading('1. Overview', 1)
    doc.add_paragraph(
        'This document describes the specifications for a new mobile application. '
        'The application will offer users real-time alerts and project management '
        'features. We aim to build a streamlined, easy-to-use platform that '
        'boosts team efficiency.'
    )

    # Features section (some overlap, some new)
    doc.add_heading('2. Core Features', 1)
    doc.add_paragraph('Key functionality includes:')

    features = [
        'User login with OAuth2 authentication',
        'Real-time push notifications for updates',
        'Calendar synchronization with multiple platforms',
        'Team collaboration tools and shared workspaces',
        'Offline mode support for uninterrupted work',
        'File attachment and sharing capabilities'
    ]

    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    # Technical Specifications (different technologies)
    doc.add_heading('3. Technology Stack', 1)
    doc.add_paragraph('Proposed technology choices:')

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Medium Grid 1 Accent 1'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology'
    hdr_cells[2].text = 'Version'

    # Data rows (some different)
    data = [
        ('Frontend', 'React Native', '0.72'),
        ('Backend', 'Express.js', '4.18'),
        ('Database', 'MongoDB', '6.0'),
        ('Cache', 'Redis', '7.0')
    ]

    for i, (comp, tech, ver) in enumerate(data, start=1):
        row = table.rows[i].cells
        row[0].text = comp
        row[1].text = tech
        row[2].text = ver

    # Project Schedule (different from doc1)
    doc.add_heading('4. Project Schedule', 1)
    doc.add_paragraph(
        'Sprint 1: Planning and setup - 2 weeks\n'
        'Sprint 2-3: Core development - 6 weeks\n'
        'Sprint 4: Feature completion - 3 weeks\n'
        'Sprint 5: Testing phase - 3 weeks\n'
        'Sprint 6: Launch preparation - 1 week'
    )

    # Additional section (not in doc1)
    doc.add_heading('5. Budget Estimates', 1)
    doc.add_paragraph(
        'Development costs: $150,000\n'
        'Infrastructure: $20,000/year\n'
        'Maintenance: $30,000/year'
    )

    doc.save('data/uploads/test_doc2.docx')
    print('✅ Created test_doc2.docx')


def create_identical_doc():
    """Create a third document identical to doc1 for testing perfect match."""
    doc = Document()

    # Same as doc1
    title = doc.add_heading('Product Requirements Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('1. Overview', 1)
    doc.add_paragraph(
        'This document outlines the requirements for the new mobile application. '
        'The app will provide users with real-time notifications and task management '
        'capabilities. Our goal is to create an intuitive, user-friendly interface '
        'that enhances productivity.'
    )

    doc.save('data/uploads/test_doc3_identical.docx')
    print('✅ Created test_doc3_identical.docx (identical to doc1)')


if __name__ == '__main__':
    print('Creating test documents...')
    create_test_doc1()
    create_test_doc2()
    create_identical_doc()
    print('\n✅ All test documents created successfully!')
    print('Documents saved in: data/uploads/')
