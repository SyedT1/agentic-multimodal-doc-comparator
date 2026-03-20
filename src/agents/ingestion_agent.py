"""
Document ingestion agent for extracting content from PDF and DOCX files.
Supports both PyMuPDF and pypdf for PDF parsing.
"""
import pdfplumber
from docx import Document
from typing import Dict, List, Any
from pathlib import Path

from agents.base_agent import BaseAgent
from models.document import RawDocument

# Try to import PyMuPDF, fallback to pypdf if not available
try:
    import fitz  # PyMuPDF
    USING_PYMUPDF = True
    print("✓ Using PyMuPDF for PDF text extraction")
except (ImportError, OSError) as e:
    print(f"⚠ PyMuPDF not available ({e}), falling back to pypdf")
    try:
        from pypdf import PdfReader
        USING_PYMUPDF = False
        print("✓ Using pypdf for PDF text extraction")
    except ImportError:
        raise ImportError(
            "Neither PyMuPDF nor pypdf is available. "
            "Install one of them: pip install PyMuPDF or pip install pypdf"
        )


class IngestionAgent(BaseAgent):
    """Agent responsible for extracting raw content from documents."""

    def __init__(self, config: Dict[str, Any] = None):
        super().__init__(config)

    def get_agent_name(self) -> str:
        return "IngestionAgent"

    async def process(self, file_path: str) -> RawDocument:
        """
        Process a document file and extract raw content.

        Args:
            file_path: Path to PDF or DOCX file

        Returns:
            RawDocument containing extracted content
        """
        file_type = self._detect_file_type(file_path)

        if file_type == "pdf":
            return await self._ingest_pdf(file_path)
        elif file_type == "docx":
            return await self._ingest_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _detect_file_type(self, file_path: str) -> str:
        """Detect file type from extension."""
        extension = Path(file_path).suffix.lower()
        if extension == ".pdf":
            return "pdf"
        elif extension in [".docx", ".doc"]:
            return "docx"
        else:
            raise ValueError(f"Unsupported file extension: {extension}")

    async def _ingest_pdf(self, file_path: str) -> RawDocument:
        """
        Extract content from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            RawDocument with extracted content
        """
        pages = []
        raw_text = ""
        raw_tables = []

        # Extract text using PyMuPDF or pypdf
        if USING_PYMUPDF:
            # Extract text using PyMuPDF
            with fitz.open(file_path) as pdf_doc:
                for page_num, page in enumerate(pdf_doc, start=1):
                    page_text = page.get_text()
                    raw_text += page_text + "\n"
                    pages.append({
                        "page_num": page_num,
                        "text": page_text
                    })
        else:
            # Extract text using pypdf
            reader = PdfReader(file_path)
            for page_num, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text() or ""
                raw_text += page_text + "\n"
                pages.append({
                    "page_num": page_num,
                    "text": page_text
                })

        # Extract tables using pdfplumber (works with both)
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                tables_on_page = page.extract_tables()
                if tables_on_page:
                    for table_idx, table in enumerate(tables_on_page):
                        if table:  # Skip empty tables
                            raw_tables.append({
                                "page_num": page_num,
                                "table_idx": table_idx,
                                "data": table
                            })

        return RawDocument(
            filename=Path(file_path).name,
            file_type="pdf",
            pages=pages,
            raw_text=raw_text.strip(),
            raw_tables=raw_tables,
            total_pages=len(pages),
            metadata={"file_path": file_path}
        )

    async def _ingest_docx(self, file_path: str) -> RawDocument:
        """
        Extract content from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            RawDocument with extracted content
        """
        doc = Document(file_path)
        pages = []
        raw_text = ""
        raw_tables = []

        # Extract text from paragraphs
        # Note: DOCX doesn't have "pages" like PDF, so we simulate page 1
        page_text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                page_text += para.text + "\n"
                raw_text += para.text + "\n"

        pages.append({
            "page_num": 1,
            "text": page_text
        })

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)

            if table_data:  # Skip empty tables
                raw_tables.append({
                    "page_num": 1,
                    "table_idx": table_idx,
                    "data": table_data
                })

        return RawDocument(
            filename=Path(file_path).name,
            file_type="docx",
            pages=pages,
            raw_text=raw_text.strip(),
            raw_tables=raw_tables,
            total_pages=1,  # DOCX treated as single page
            metadata={"file_path": file_path}
        )
